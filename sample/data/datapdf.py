# /// script
# requires-python = ">=3.10"
# dependencies = ["pymupdf", "python-docx"]
# ///

"""
法令遵循處資料處理腳本
用法：uv run datapdf.py

從 法令遵循處/ 資料夾讀取資料，輸出至同層的 data/ 資料夾：
  裁罰.json          ← 重大裁罰、非重大裁罰文件
  法規.json          ← 法規修訂文件
  主管法規資料集.jsonl ← 各局主管法規 JSONL 合併（加入機構名稱）
  全國法規資料庫.jsonl ← 直接複製
"""

import os
import json
import shutil

# ── 路徑設定 ─────────────────────────────────────────────────────
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
BASE         = os.path.join(SCRIPT_DIR, "法令遵循處", "法令遵循處")
PENALTY_DIR  = os.path.join(BASE, "重大裁罰、非重大裁罰、法規修訂", "裁罰")
REG_DIR      = os.path.join(BASE, "重大裁罰、非重大裁罰、法規修訂", "法規")
FSC_DIR      = os.path.join(BASE, "主管法規資料集")
NAT_SRC      = os.path.join(FSC_DIR, "全國法規資料庫.jsonl")

OUT_PENALTY  = os.path.join(SCRIPT_DIR, "裁罰.json")
OUT_REG      = os.path.join(SCRIPT_DIR, "法規.json")
OUT_FSC      = os.path.join(SCRIPT_DIR, "主管法規資料集.jsonl")
OUT_NAT      = os.path.join(SCRIPT_DIR, "全國法規資料庫.jsonl")
# ─────────────────────────────────────────────────────────────────


def read_txt(path):
    for enc in ("utf-8", "big5", "cp950"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeDecodeError(f"無法解碼 {path}")


def read_pdf(path):
    import fitz
    doc = fitz.open(path)
    text = "\n".join(page.get_text("text") for page in doc)
    doc.close()
    return text


def read_docx(path):
    import docx
    doc = docx.Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_file(path):
    ext = path.rsplit(".", 1)[-1].lower()
    if ext == "txt":
        return read_txt(path)
    elif ext == "pdf":
        return read_pdf(path)
    elif ext == "docx":
        return read_docx(path)
    return None


def clean_lines(raw):
    return [l.strip() for l in raw.split("\n") if l.strip()]


# ── 1. 裁罰 ──────────────────────────────────────────────────────
def process_penalties():
    if not os.path.exists(PENALTY_DIR):
        print(f"[SKIP] 找不到裁罰資料夾：{PENALTY_DIR}")
        return

    records, ok, err = [], 0, 0
    print("處理裁罰...")

    for root, _, files in os.walk(PENALTY_DIR):
        for fname in files:
            if fname.startswith("."):
                continue
            fpath = os.path.join(root, fname)
            ext   = fname.rsplit(".", 1)[-1].lower()
            if ext not in ("txt", "pdf", "docx"):
                continue

            rel   = os.path.relpath(root, PENALTY_DIR)
            parts = rel.split(os.sep) if rel != "." else []

            category    = parts[0] if len(parts) > 0 else "未分類"
            institution = parts[1] if len(parts) > 1 else "未分類"
            time_str    = parts[2] if len(parts) > 2 else "未知時間"

            # 標題取檔名（去掉尾端的 _YYYY-MM-DD）
            title = os.path.splitext(fname)[0]
            if "_" in title:
                stem, maybe_date = title.rsplit("_", 1)
                if len(maybe_date) == 10 and maybe_date.count("-") == 2:
                    title = stem

            try:
                raw = read_file(fpath)
                if not raw:
                    continue
                lines = clean_lines(raw)
                content = "\n".join(lines)
                records.append({
                    "資料類別": category,
                    "機構名稱": institution,
                    "標題":     title,
                    "時間":     time_str,
                    "內文":     content,
                })
                ok += 1
            except Exception as e:
                print(f"  [ERR] {fname}: {e}")
                err += 1

    with open(OUT_PENALTY, "w", encoding="utf-8") as f:
        json.dump(records, f, ensure_ascii=False, indent=2)

    print(f"  完成：{ok} 筆，失敗 {err} 筆 → {OUT_PENALTY}")


# ── 2. 法規 ──────────────────────────────────────────────────────
def process_regulations():
    if not os.path.exists(REG_DIR):
        print(f"[SKIP] 找不到法規資料夾：{REG_DIR}")
        return

    records, ok, err = [], 0, 0
    print("處理法規...")

    for root, _, files in os.walk(REG_DIR):
        for fname in files:
            if fname.startswith("."):
                continue
            fpath = os.path.join(root, fname)
            ext   = fname.rsplit(".", 1)[-1].lower()
            if ext not in ("txt", "pdf", "docx"):
                continue

            rel   = os.path.relpath(root, REG_DIR)
            parts = rel.split(os.sep) if rel != "." else []

            category = parts[0] if len(parts) > 0 else "未分類"   # 機構 (金管會…)
            time_str = parts[1] if len(parts) > 1 else "未知時間"  # 日期資料夾

            title = os.path.splitext(fname)[0]
            if "_" in title:
                stem, maybe_date = title.rsplit("_", 1)
                if len(maybe_date) == 10 and maybe_date.count("-") == 2:
                    title = stem

            try:
                raw = read_file(fpath)
                if not raw:
                    continue
                lines = clean_lines(raw)
                content = "\n".join(lines)
                records.append({
                    "資料類別": category,
                    "標題":     title,
                    "時間":     time_str,
                    "內文":     content,
                })
                ok += 1
            except Exception as e:
                print(f"  [ERR] {fname}: {e}")
                err += 1

    with open(OUT_REG, "w", encoding="utf-8") as f:
        json.dump(records, f, ensure_ascii=False, indent=2)

    print(f"  完成：{ok} 筆，失敗 {err} 筆 → {OUT_REG}")


# ── 3. 主管法規資料集 JSONL 合併 ─────────────────────────────────
def process_fsc_jsonl():
    if not os.path.exists(FSC_DIR):
        print(f"[SKIP] 找不到主管法規資料集資料夾：{FSC_DIR}")
        return

    total = 0
    print("處理主管法規資料集...")

    with open(OUT_FSC, "w", encoding="utf-8") as out:
        for fname in sorted(os.listdir(FSC_DIR)):
            if not fname.endswith(".jsonl") or fname == "全國法規資料庫.jsonl":
                continue

            # 從檔名推斷機構名稱：「銀行局主管法規資料集.jsonl」→「銀行局」
            inst = fname.replace("主管法規資料集.jsonl", "").strip()

            fpath = os.path.join(FSC_DIR, fname)
            count = 0
            with open(fpath, encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        data = json.loads(line)
                        if "機構名稱" not in data:
                            data = {"機構名稱": inst, **data}
                        out.write(json.dumps(data, ensure_ascii=False) + "\n")
                        count += 1
                    except json.JSONDecodeError as e:
                        print(f"  [ERR] {fname} 解析失敗：{e}")
            print(f"  {fname}: {count} 筆")
            total += count

    print(f"  合併完成：共 {total} 筆 → {OUT_FSC}")


# ── 4. 全國法規資料庫 複製 ────────────────────────────────────────
def copy_national_laws():
    if not os.path.exists(NAT_SRC):
        print(f"[SKIP] 找不到全國法規資料庫：{NAT_SRC}")
        return
    shutil.copy2(NAT_SRC, OUT_NAT)
    count = sum(1 for l in open(OUT_NAT, encoding="utf-8") if l.strip())
    print(f"全國法規資料庫：{count} 筆 → {OUT_NAT}")


# ── 主程式 ────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 50)
    process_penalties()
    print()
    process_regulations()
    print()
    process_fsc_jsonl()
    print()
    copy_national_laws()
    print("=" * 50)
    print("全部完成。")
